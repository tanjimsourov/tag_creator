from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Advikon_Tag_Creator_Source_Strategy_Report.docx")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_borders(table, color: str = "D9E2F3") -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tbl_pr.append(borders)


def cell_text(cell, text: str, bold: bool = False, size: float = 7.3, color: str = "2D2D2D") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def para(doc: Document, text: str, size: float = 7.7, bold: bool = False, color: str = "2D2D2D") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string("1F4E79")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(7.7)

    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header.columns[0].width = Inches(5.0)
    header.columns[1].width = Inches(2.2)
    set_borders(header, "1F4E79")
    for c in header.rows[0].cells:
        shade_cell(c, "1F4E79")
    cell_text(header.rows[0].cells[0], "Advikon | Local Hero Tag Creator", True, 11, "FFFFFF")
    cell_text(header.rows[0].cells[1], "Source Strategy | 29 Jun 2026", True, 7.4, "FFFFFF")
    header.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Free Metadata First + Paid AI Only for Gaps")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("1F4E79")

    para(
        doc,
        "Recommendation: keep tag_creator as the Local Hero metadata pipeline. Use free/catalog sources first, then call SONOTELLER only for missing or advanced playlist tags. This reduces paid requests while keeping the dataset rich enough for dynamic playlists.",
        7.9,
    )

    heading(doc, "Source Roles")
    table = doc.add_table(rows=4, cols=3)
    table.autofit = False
    table.columns[0].width = Inches(1.25)
    table.columns[1].width = Inches(3.0)
    table.columns[2].width = Inches(3.0)
    set_borders(table)
    for i, h in enumerate(["Source", "Best For", "Limit / Use Rule"]):
        shade_cell(table.rows[0].cells[i], "1F4E79")
        cell_text(table.rows[0].cells[i], h, True, 7.5, "FFFFFF")

    rows = [
        (
            "Free/catalog APIs",
            "MusicBrainz/AcoustID identify recordings; Spotify gives catalog metadata; Last.fm gives tag/genre signals; Discogs gives label/catalog/style; Cover Art Archive gives artwork.",
            "Use first for title, artist, album, release date, ISRC, label, cover art, and genre hints. Low cost, but weaker for mood/subgenre/BPM/key.",
        ),
        (
            "SONOTELLER paid AI",
            "AI analysis for genre, subgenre, moods, instruments, BPM, key, language, themes, sections and music/lyrics summaries.",
            "Use only when required advanced tags are missing or confidence is low. Needs RapidAPI key and file URL or confirmed upload endpoint.",
        ),
        (
            "tag_creator",
            "Batch orchestration, CSV cache/state, conflict detection, field mapping, reports, and Local Hero integration.",
            "This is the permanent automation layer; paid APIs are provider plugins inside it, not replacements.",
        ),
    ]
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, text in enumerate(row):
            if c_idx == 0:
                shade_cell(table.rows[r_idx].cells[c_idx], "EAF2F8")
                cell_text(table.rows[r_idx].cells[c_idx], text, True, 7.3, "1F4E79")
            else:
                cell_text(table.rows[r_idx].cells[c_idx], text, False, 7.0)

    heading(doc, "Cost Control Plan")
    para(
        doc,
        "Example: if 5,000 files exist, tag_creator first enriches all files with free sources. If 3,000 become complete enough, only 2,000 are sent to SONOTELLER. This can save thousands of paid requests.",
        7.7,
    )
    price = doc.add_table(rows=2, cols=4)
    price.autofit = False
    for col in price.columns:
        col.width = Inches(1.8)
    set_borders(price)
    for i, h in enumerate(["Pro", "Ultra", "Mega", "Decision"]):
        shade_cell(price.rows[0].cells[i], "1F4E79")
        cell_text(price.rows[0].cells[i], h, True, 7.4, "FFFFFF")
    price_values = [
        "$7.95 / 50 req. (~$0.16/file)",
        "$69.95 / 500 req. (~$0.14/file)",
        "$599.95 / 5,000 req. (~$0.12/file)",
        "For thousands: Mega or custom/bulk pricing after a small pilot.",
    ]
    for i, val in enumerate(price_values):
        cell_text(price.rows[1].cells[i], val, False, 7.0)

    heading(doc, "Implementation Path")
    para(
        doc,
        "1) Finish free-source adapters and missing-tag detector. 2) Buy a small SONOTELLER/RapidAPI plan for pilot. 3) Test 20-50 owned files. 4) Confirm request count, endpoint, URL/upload requirement, and response fields. 5) Run full DB in batches and export Local Hero CSV/DB-ready metadata.",
        7.7,
    )

    heading(doc, "References")
    refs = (
        "SONOTELLER pricing: https://rapidapi.com/sonoteller1-sonoteller-default/api/sonoteller-ai1/pricing | "
        "SONOTELLER capabilities: https://sonoteller.ai/ | "
        "Spotify Web API: https://developer.spotify.com/documentation/web-api | "
        "MusicBrainz API: https://musicbrainz.org/doc/MusicBrainz_API | "
        "Last.fm API: https://www.last.fm/api | "
        "Discogs API: https://www.discogs.com/developers | "
        "Cover Art Archive API: https://musicbrainz.org/doc/Cover_Art_Archive/API | "
        "AcoustID API: https://acoustid.org/webservice"
    )
    para(doc, refs, 5.8, False, "555555")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Created by: Tanjim Sourov | tanjim@advikon.eu | Soft Dev, Advikon")
    run.font.name = "Aptos"
    run.font.size = Pt(7)
    run.italic = True
    run.font.color.rgb = RGBColor.from_string("666666")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT.resolve())
