from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Advikon_Tag_Creator_Local_Hero_Report.docx")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9E2F3")
        borders.append(tag)
    tbl_pr.append(borders)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)


def add_body(doc: Document, text: str, size: int = 8) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(45, 45, 45)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(8)

    # Header band
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(2.4)
    set_borders(table)
    for cell in table.rows[0].cells:
        shade_cell(cell, "1F4E79")
    left, right = table.rows[0].cells
    set_cell_text(left, "Advikon | Local Hero Metadata Automation", True, "FFFFFF", 12)
    set_cell_text(right, "Report Date: 26 June 2026", True, "FFFFFF", 8)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    title = p.add_run("Tag Creator Tool: Purpose, Status, and Local Hero Role")
    title.bold = True
    title.font.size = Pt(15)
    title.font.color.rgb = RGBColor(31, 78, 121)

    add_body(
        doc,
        "This report explains why the custom tag_creator tool is needed for Local Hero task #4: screening the music database with correct and richer tags. The goal is not a one-time CSV export, but a reusable automation component that keeps the music dataset ready for dynamic playlist generation.",
        8,
    )

    add_heading(doc, "Current Position")
    status = doc.add_table(rows=4, cols=2)
    status.autofit = False
    status.columns[0].width = Inches(1.7)
    status.columns[1].width = Inches(5.5)
    set_borders(status)
    rows = [
        ("Basic build", "Initial Python tool is already built: scan files, read MP3/MP4 tags, export CSV, use CSV cache/state, and run safe dry-runs."),
        ("Why not only Music Tag", "Music Tag is useful for basic metadata and artwork, but it does not reliably provide the larger business tags needed for playlist automation."),
        ("Pro target", "Combine multiple metadata sources, verify conflicts, enrich missing fields, write safe tags, and produce reports for thousands of tracks."),
        ("Timeline", "Basic version took 1-2 days. A production-ready version needs more than 1 week, realistically 1-2 weeks for robust provider integration and testing."),
    ]
    for idx, (label, value) in enumerate(rows):
        c0, c1 = status.rows[idx].cells
        shade_cell(c0, "EAF2F8")
        set_cell_text(c0, label, True, "1F4E79", 8)
        set_cell_text(c1, value, False, "2D2D2D", 8)

    add_heading(doc, "Music Tag vs. tag_creator")
    comp = doc.add_table(rows=4, cols=3)
    comp.autofit = False
    comp.columns[0].width = Inches(1.8)
    comp.columns[1].width = Inches(2.4)
    comp.columns[2].width = Inches(3.0)
    set_borders(comp)
    headers = ["Area", "Music Tag paid tool", "Custom tag_creator"]
    for i, text in enumerate(headers):
        shade_cell(comp.rows[0].cells[i], "1F4E79")
        set_cell_text(comp.rows[0].cells[i], text, True, "FFFFFF", 8)
    comp_rows = [
        ("Main job", "Updates common tags and artwork inside files.", "Builds verified metadata dataset for Local Hero automation."),
        ("Tag depth", "Mostly title, artist, album, year, track, artwork.", "Adds richer fields such as genre, language, country, mood, season, weather/event suitability, confidence, and missing-tag status."),
        ("Automation role", "Manual/semi-manual app workflow.", "Repeatable backend process for new music batches and playlist intelligence."),
    ]
    for row_idx, row in enumerate(comp_rows, start=1):
        for col_idx, text in enumerate(row):
            if col_idx == 0:
                shade_cell(comp.rows[row_idx].cells[col_idx], "F3F6FA")
                set_cell_text(comp.rows[row_idx].cells[col_idx], text, True, "1F4E79", 8)
            else:
                set_cell_text(comp.rows[row_idx].cells[col_idx], text, False, "2D2D2D", 7)

    add_heading(doc, "Recommended Direction")
    add_body(
        doc,
        "Continue building tag_creator as part of the Local Hero automation layer. Use Music Tag only as a helper for basic file metadata when needed. tag_creator should become the controlled enrichment process that runs whenever new songs are added, then outputs clean CSV data for playlist generation.",
        8,
    )

    callout = doc.add_table(rows=1, cols=1)
    set_borders(callout)
    shade_cell(callout.rows[0].cells[0], "FFF2CC")
    set_cell_text(
        callout.rows[0].cells[0],
        "Conclusion: tag_creator is directly related to task #4 and becomes the metadata foundation for later dynamic playlist automation. It is not just one-time CSV processing; it should continue running as part of the Local Hero music workflow.",
        True,
        "7F6000",
        8,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Created by: Tanjim Sourov | tanjim@advikon.eu | Soft Dev, Advikon")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 90, 90)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT.resolve())
