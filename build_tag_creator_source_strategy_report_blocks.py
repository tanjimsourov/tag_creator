from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Advikon_Tag_Creator_Source_Strategy_Report.docx")


def set_paragraph_block(paragraph, fill: str, border: str = "D9E2F3") -> None:
    p_pr = paragraph._p.get_or_add_pPr()

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)

    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "4")
        tag.set(qn("w:color"), border)
        borders.append(tag)
    p_pr.append(borders)


def add_run(paragraph, text: str, *, bold: bool = False, size: float = 8, color: str = "2D2D2D") -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_block(
    doc: Document,
    title: str,
    body: str,
    *,
    fill: str = "F7FAFC",
    border: str = "D9E2F3",
    title_color: str = "1F4E79",
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_block(p, fill, border)
    add_run(p, title, bold=True, size=8.5, color=title_color)
    add_run(p, "\n" + body, size=7.4, color="2D2D2D")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(7.6)

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(5)
    header.paragraph_format.line_spacing = 1.0
    set_paragraph_block(header, "1F4E79", "1F4E79")
    add_run(header, "Advikon | Local Hero Tag Creator\n", bold=True, size=12, color="FFFFFF")
    add_run(header, "Free Metadata First + Paid AI Only for Gaps | 29 Jun 2026", size=7.6, color="DDEBF7")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(4)
    intro.paragraph_format.line_spacing = 1.0
    add_run(
        intro,
        "Recommendation: keep tag_creator as the permanent Local Hero metadata pipeline. Use free/catalog sources first, then call SONOTELLER only for missing or advanced playlist tags. This reduces paid requests while keeping the dataset rich enough for dynamic playlists.",
        size=7.8,
    )

    add_block(
        doc,
        "Source Strategy",
        "Free/catalog APIs identify what the song is: MusicBrainz/AcoustID for recording identity, Spotify for catalog metadata, Last.fm for tag/genre signals, Discogs for label/catalog/style, and Cover Art Archive for artwork. SONOTELLER adds what the song feels like: genre, subgenre, moods, instruments, BPM, key, language, themes and sections.",
        fill="EAF2F8",
        border="9ECAE1",
    )

    add_block(
        doc,
        "Why tag_creator is still required",
        "Paid APIs are data sources, not the full Local Hero workflow. tag_creator scans thousands of MP3/MP4 files, runs providers in batches, avoids rate-limit failures, compares conflicting metadata, stores CSV cache/state, maps fields to Local Hero needs, and exports DB-ready reports.",
        fill="F7FAFC",
        border="D9E2F3",
    )

    add_block(
        doc,
        "Cost Control",
        "Example: if 5,000 files exist, tag_creator first enriches all files with free sources. If 3,000 become complete enough, only 2,000 are sent to SONOTELLER. Visible RapidAPI pricing: Pro $7.95 / 50 requests (~$0.16/file), Ultra $69.95 / 500 (~$0.14/file), Mega $599.95 / 5,000 (~$0.12/file). For thousands, Mega or custom/bulk pricing is realistic.",
        fill="FFF2CC",
        border="D6B656",
        title_color="7F6000",
    )

    add_block(
        doc,
        "Implementation Path",
        "1) Finish free-source adapters and missing-tag detector. 2) Buy a small SONOTELLER/RapidAPI plan for pilot. 3) Test 20-50 owned files. 4) Confirm request count, endpoint, URL/upload requirement and response fields. 5) Run full DB in batches and export Local Hero CSV/DB-ready metadata.",
        fill="E2F0D9",
        border="A9D18E",
        title_color="548235",
    )

    refs = doc.add_paragraph()
    refs.paragraph_format.space_before = Pt(2)
    refs.paragraph_format.space_after = Pt(3)
    refs.paragraph_format.line_spacing = 1.0
    set_paragraph_block(refs, "F2F2F2", "D9D9D9")
    add_run(refs, "References: ", bold=True, size=6.6, color="555555")
    add_run(
        refs,
        "SONOTELLER pricing https://rapidapi.com/sonoteller1-sonoteller-default/api/sonoteller-ai1/pricing | "
        "Capabilities https://sonoteller.ai/ | Spotify https://developer.spotify.com/documentation/web-api | "
        "MusicBrainz https://musicbrainz.org/doc/MusicBrainz_API | Last.fm https://www.last.fm/api | "
        "Discogs https://www.discogs.com/developers | Cover Art Archive https://musicbrainz.org/doc/Cover_Art_Archive/API | "
        "AcoustID https://acoustid.org/webservice",
        size=5.8,
        color="555555",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(2)
    footer.paragraph_format.space_after = Pt(0)
    add_run(footer, "Created by: Tanjim Sourov | tanjim@advikon.eu | Soft Dev, Advikon", size=7, color="666666")
    footer.runs[0].italic = True

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT.resolve())
